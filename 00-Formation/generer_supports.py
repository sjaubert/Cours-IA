"""
Génère les 2 supports imprimés pour la formation IA AFPI :
  1. Fiche_Apprenant_Carnet_de_Prompts.docx
  2. Fiche_Apprenant_CheatSheet_Prompt.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

LOGO = os.path.join(os.path.dirname(__file__), '..', 'logo_uimm_placeholder.jpg')
OUT_DIR = os.path.dirname(__file__)

BLEU = RGBColor(0x1A, 0x37, 0x6C)   # bleu UIMM
GRIS = RGBColor(0x60, 0x60, 0x60)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)

# ──────────────────────────────────────────────
# Helpers
# ──────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_header(doc, title, subtitle=""):
    """En-tête avec logo + titre sur fond bleu."""
    # Tableau en-tête : logo | titre
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.columns[0].width = Cm(3.5)
    tbl.columns[1].width = Cm(13.5)
    cell_logo = tbl.cell(0, 0)
    cell_title = tbl.cell(0, 1)
    set_cell_bg(cell_logo, '1A376C')
    set_cell_bg(cell_title, '1A376C')
    try:
        p = cell_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(LOGO, width=Cm(2.8))
    except Exception:
        pass
    p2 = cell_title.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run2 = p2.add_run("Pôle Formation UIMM - CVDL")
    run2.font.color.rgb = BLANC
    run2.font.bold = True
    run2.font.size = Pt(10)
    p2.add_run("\n")
    run3 = p2.add_run(title)
    run3.font.color.rgb = BLANC
    run3.font.bold = True
    run3.font.size = Pt(14)
    if subtitle:
        p2.add_run("\n")
        run4 = p2.add_run(subtitle)
        run4.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
        run4.font.size = Pt(9)
    doc.add_paragraph()

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLEU
    # Ligne de soulignement via bordure basse
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A376C')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_lecon(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Lecon : ")
    run.font.bold = True
    run.font.color.rgb = BLEU
    run.font.size = Pt(10)
    run2 = p.add_run(texte)
    run2.font.italic = True
    run2.font.size = Pt(10)
    run2.font.color.rgb = GRIS

def add_prompt_box(doc, label, prompt_text):
    """Boite gris clair pour un prompt."""
    p_label = doc.add_paragraph()
    p_label.paragraph_format.space_after = Pt(2)
    r = p_label.add_run(label)
    r.font.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = BLEU

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'EEF2F7')
    p = cell.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(prompt_text)
    run.font.size = Pt(9)
    run.font.name = 'Courier New'
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_obs(doc, texte):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Observation : ")
    r.font.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xB8, 0x6A, 0x00)
    r2 = p.add_run(texte)
    r2.font.size = Pt(9)
    r2.font.color.rgb = GRIS

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Pôle Formation UIMM - CVDL  |  Formateur : S. JAUBERT  |  Reproduction interdite sans accord")
    run.font.size = Pt(8)
    run.font.color.rgb = GRIS

def body_normal(doc, text, bold=False, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.bold = bold


# ══════════════════════════════════════════════
# DOCUMENT 1 : Carnet de Prompts
# ══════════════════════════════════════════════

def build_carnet():
    doc = Document()
    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    add_header(doc, "Carnet de Prompts", "Formation Intelligence Artificielle — Ateliers Pratiques")
    add_footer(doc)

    body_normal(doc,
        "Ce carnet rassemble les exercices pratiques de la formation. "
        "Copiez-collez les prompts dans l'outil IA mis à votre disposition "
        "(NotebookLM ou Antigravity selon l'atelier).")
    doc.add_paragraph()

    # ── Atelier 1
    add_section_title(doc, "Atelier 1 — La Puissance du Contexte")
    add_lecon(doc,
        "Plus vous donnez de contexte à l'IA (rôle, objectif, contraintes), "
        "plus sa réponse sera pertinente et immédiatement utilisable.")

    add_prompt_box(doc, "Test A — Requête générique (sans contexte) :",
        "Rédige un message pour dire à l'équipe que la réunion de lundi est repoussée à mercredi matin.")
    add_obs(doc,
        "Notez la froideur et le manque de personnalisation. "
        "L'IA a 'deviné' sans savoir qui parle à qui.")

    add_prompt_box(doc, "Test B — Requête contextualisée (Prompt structuré) :",
        "Tu es un manager bienveillant dans une grande entreprise industrielle.\n"
        "Ton objectif est d'informer ton équipe de production (15 personnes) que la réunion\n"
        "hebdomadaire de lundi prochain à 9h est reportée à mercredi 10h en raison d'une panne\n"
        "inopinée sur la ligne d'assemblage 4 qu'il faut régler en urgence.\n"
        "Ta réponse doit être courte, rassurante et se terminer par un mot d'encouragement.")
    add_obs(doc,
        "L'IA s'adapte radicalement. En réduisant son espace de probabilités, "
        "vous obtenez un résultat directement exploitable.")

    # ── Atelier 2
    add_section_title(doc, "Atelier 2 — La Température (Factualité vs Créativité)")
    add_lecon(doc,
        "La température définit le niveau de créativité du modèle. "
        "Vous pouvez la 'simuler' par le texte si l'interface n'a pas de réglage dédié.")

    add_prompt_box(doc, "Test A — Température basse (analyse factuelle) :",
        "Je souhaite que tes réponses soient le plus factuelles, concises et directes possibles.\n"
        "N'invente rien. Explique-moi le fonctionnement d'un moteur à explosion en 3 phrases simples.")

    add_prompt_box(doc, "Test B — Température haute (brainstorming créatif) :",
        "Je souhaite que tu répondes de la manière la plus créative et décalée possible.\n"
        "Explique-moi le fonctionnement d'un moteur à explosion\n"
        "comme si tu racontais une épopée chevaleresque.")
    add_obs(doc,
        "Le vocabulaire est radicalement différent alors que la question est identique. "
        "Choisissez la 'température' selon votre besoin : analyse ou idéation.")

    # ── Atelier 3
    add_section_title(doc, "Atelier 3 — Induire une Hallucination")
    add_lecon(doc,
        "Un LLM n'est pas un moteur de recherche. S'il ne sait pas, "
        "il invente une réponse plausible avec beaucoup d'aplomb.")

    add_prompt_box(doc, "Test — Le Piège :",
        "Fais-moi un résumé concis du traité de 'Neuville-sur-Vance' qui a mis fin\n"
        "au grand conflit syndical de la métallurgie en 1984 dans la région Grand-Est.\n"
        "Parle-moi des 3 grandes mesures phares de ce traité.")
    add_obs(doc,
        "Ce conflit n'existe pas. Pourtant l'IA va inventer des noms, dates et mesures plausibles. "
        "C'est pourquoi le RAG (fournir ses propres documents) est indispensable en contexte professionnel.")

    # ── Atelier NotebookLM
    add_section_title(doc, "Atelier 4 — NotebookLM : Transformer ses documents en assistant")
    add_lecon(doc,
        "NotebookLM est un outil RAG 'zero code'. Il ne repond qu'a partir de VOS documents. "
        "Aucun risque d'hallucination sur vos sources.")

    p_prep = doc.add_paragraph()
    p_prep.paragraph_format.space_after = Pt(4)
    r_prep = p_prep.add_run("Avant l'atelier — Preparez votre document :")
    r_prep.font.bold = True
    r_prep.font.size = Pt(10)
    r_prep.font.color.rgb = BLEU

    tbl_doc = doc.add_table(rows=1, cols=1)
    tbl_doc.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_doc = tbl_doc.cell(0, 0)
    set_cell_bg(cell_doc, 'FFF8E1')
    p_doc = cell_doc.paragraphs[0]
    p_doc.paragraph_format.left_indent = Cm(0.3)
    p_doc.paragraph_format.space_before = Pt(4)
    p_doc.paragraph_format.space_after = Pt(4)
    consigne = (
        "Ouvrez sur votre poste un fichier PDF ou Word que vous utilisez dans votre travail.\n"
        "Par exemple :\n"
        "  - une procedure ou instruction de travail\n"
        "  - un compte-rendu de reunion\n"
        "  - une fiche produit ou notice technique\n"
        "  - un reglement interieur ou une note RH\n"
        "  - un programme de formation ou un referentiel de competences\n"
        "  - un rapport d'audit ou un bilan d'activite\n"
        "  - une fiche de poste\n\n"
        "Ce document restera confidentiel : il ne sera utilise que par vous, sur votre poste."
    )
    r_doc = p_doc.add_run(consigne)
    r_doc.font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    add_prompt_box(doc, "Etape 1 — Creer un notebook et importer votre document :",
        "1. Aller sur notebooklm.google.com\n"
        "2. Cliquer sur 'Nouveau notebook'\n"
        "3. Importer le fichier que vous avez prepare\n"
        "4. Attendre l'indexation (30 secondes environ)")

    add_prompt_box(doc, "Etape 2 — Prompts à tester sur votre document :",
        "Résume ce document en 5 points clés.\n\n"
        "Génère 5 questions de quiz avec les bonnes réponses basées sur ce document.\n\n"
        "Crée un tableau comparatif des avantages et inconvénients mentionnés.\n\n"
        "Rédige une fiche de synthèse d'une page pour un public non-expert.")
    add_obs(doc,
        "Toutes les réponses sont ancrées dans VOS documents. "
        "Cliquez sur les sources pour vérifier. NotebookLM peut aussi générer un podcast audio !")

    # ── Aide-mémoire
    add_section_title(doc, "Aide-Memoire — Structure universelle d'un prompt")
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Table Grid'
    headers = [("ROLE", "Agis en tant que spécialiste qualité industrielle..."),
               ("TACHE", "Analyse ces trois rapports d'incidents..."),
               ("CONTEXTE", "Notre entreprise fabrique des moteurs et rencontre des défauts sur l'axe X."),
               ("FORMAT", "Organise ta réponse en tableau comparatif, sans intro."),
               ("CONTRAINTE", "Utilise uniquement les données fournies. Si tu ne sais pas, dis-le.")]
    for i, (k, v) in enumerate(headers):
        set_cell_bg(tbl.cell(i, 0), '1A376C')
        r = tbl.cell(i, 0).paragraphs[0].add_run(k)
        r.font.bold = True
        r.font.color.rgb = BLANC
        r.font.size = Pt(10)
        rv = tbl.cell(i, 1).paragraphs[0].add_run(v)
        rv.font.size = Pt(9)

    doc.add_paragraph()
    body_normal(doc,
        "Espace notes personnelles — Mon prompt metier :", bold=True)
    for _ in range(4):
        p = doc.add_paragraph("_" * 90)
        p.paragraph_format.space_after = Pt(10)

    path = os.path.join(OUT_DIR, 'Fiche_Apprenant_Carnet_de_Prompts.docx')
    doc.save(path)
    print(f"[OK] {path}")


# ══════════════════════════════════════════════
# DOCUMENT 2 : Cheat Sheet Prompt Engineering
# ══════════════════════════════════════════════

def build_cheatsheet():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    add_header(doc, "Cheat Sheet — Prompt Engineering",
               "Les techniques essentielles pour maitriser l'IA generative")
    add_footer(doc)
    add_section_title(doc, "Les frameworks de prompt")

    # Tableau AUTOMAT vs CO-STAR
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes colonnes
    for i, h in enumerate(["Framework AUTOMAT", "Framework CO-STAR"]):
        set_cell_bg(tbl.cell(0, i), '1A376C')
        r = tbl.cell(0, i).paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = BLANC
        r.font.size = Pt(11)

    automat = [
        ("A", "Act as...", "Définir le rôle / persona de l'IA"),
        ("U", "User & Audience", "Qui est l'interlocuteur ?"),
        ("T", "Targeted Action", "Quelle action précise ?"),
        ("O", "Output Definition", "Quel format de sortie ?"),
        ("M", "Mode / Tonalité", "Quel style de réponse ?"),
        ("A", "Atypical Cases", "Comment gérer les cas limites ?"),
        ("T", "Topic Whitelisting", "Quels sujets sont autorisés ?"),
    ]
    costar = [
        ("C", "Context", "Contexte et arrière-plan"),
        ("O", "Objective", "Objectif clair à atteindre"),
        ("S", "Style & Tone", "Style et tonalité souhaitée"),
        ("T", "Target Audience", "Public cible visé"),
        ("A", "Action", "Action concrète demandée"),
        ("R", "Response Format", "Format de la réponse"),
        ("", "", ""),
    ]
    for i, ((la, lkw, ldesc), (ra, rkw, rdesc)) in enumerate(zip(automat, costar)):
        row = tbl.add_row()
        # Gauche
        cl = row.cells[0]
        p = cl.paragraphs[0]
        r1 = p.add_run(f"{la} — {lkw} : ")
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = BLEU
        r1b = p.add_run(ldesc)
        r1b.font.size = Pt(9)
        # Droite
        cr = row.cells[1]
        p2 = cr.paragraphs[0]
        r2 = p2.add_run(f"{ra} — {rkw} : " if ra else "")
        r2.font.bold = True
        r2.font.size = Pt(9)
        r2.font.color.rgb = BLEU
        r2b = p2.add_run(rdesc)
        r2b.font.size = Pt(9)

    doc.add_paragraph()
    add_section_title(doc, "Les 5 techniques fondamentales")

    techniques = [
        ("1. Zero-Shot",
         "Poser une question directe sans exemple.\n"
         "Ex : « Résume ce texte en 3 points »"),
        ("2. Few-Shot (exemples)",
         "Montrer 2-3 exemples avant la vraie question.\n"
         "Ex : Entree: 'chat' → Sortie: 'animal domestique'\n"
         "     Entree: 'Ford' → Sortie: ?"),
        ("3. Chain-of-Thought (etape par etape)",
         "Forcer l'IA à raisonner avant de répondre.\n"
         "Ex : « Résous ce problème étape par étape, "
         "montre ton raisonnement avant de donner la réponse. »"),
        ("4. RAG (tes propres documents)",
         "Fournir un document source pour ancrer les réponses.\n"
         "Ex avec NotebookLM : importer son PDF puis questionner."),
        ("5. Role Prompting",
         "Attribuer un rôle expert à l'IA.\n"
         "Ex : « Tu es un expert en maintenance industrielle "
         "avec 20 ans d'expérience... »"),
    ]

    for titre, desc in techniques:
        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl2.columns[0].width = Cm(4)
        tbl2.columns[1].width = Cm(13)
        set_cell_bg(tbl2.cell(0, 0), '1A376C')
        rk = tbl2.cell(0, 0).paragraphs[0].add_run(titre)
        rk.font.bold = True
        rk.font.color.rgb = BLANC
        rk.font.size = Pt(9)
        tbl2.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(tbl2.cell(0, 1), 'EEF2F7')
        rv = tbl2.cell(0, 1).paragraphs[0].add_run(desc)
        rv.font.size = Pt(9)
        tbl2.cell(0, 1).paragraphs[0].paragraph_format.left_indent = Cm(0.3)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)

    add_section_title(doc, "Les 5 concepts cles a connaitre")
    concepts = [
        ("Token", "L'IA ne lit pas des mots mais des syllabes (tokens). 1 mot ≈ 1,3 tokens."),
        ("Fenetre de contexte", "La memoire a court terme de l'IA. Au-dela, elle 'oublie' le debut."),
        ("Temperature", "Curseur creativite : basse = factuel/precis, haute = creatif/risque."),
        ("Hallucination", "L'IA invente une reponse plausible quand elle ne sait pas. Toujours verifier."),
        ("RAG", "Retrieval-Augmented Generation : l'IA repond a partir de VOS documents. "
                "Elimine les hallucinations sur vos sources."),
    ]
    tbl3 = doc.add_table(rows=len(concepts)+1, cols=2)
    tbl3.style = 'Table Grid'
    set_cell_bg(tbl3.cell(0, 0), '1A376C')
    set_cell_bg(tbl3.cell(0, 1), '1A376C')
    for col, h in enumerate(["Concept", "Explication rapide"]):
        r = tbl3.cell(0, col).paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = BLANC
        r.font.size = Pt(10)
    for i, (c, d) in enumerate(concepts, 1):
        if i % 2 == 0:
            set_cell_bg(tbl3.cell(i, 0), 'EEF2F7')
            set_cell_bg(tbl3.cell(i, 1), 'EEF2F7')
        rc = tbl3.cell(i, 0).paragraphs[0].add_run(c)
        rc.font.bold = True
        rc.font.size = Pt(9)
        rc.font.color.rgb = BLEU
        rd = tbl3.cell(i, 1).paragraphs[0].add_run(d)
        rd.font.size = Pt(9)

    doc.add_paragraph()
    add_section_title(doc, "Erreurs classiques a eviter")
    erreurs = [
        "Trop vague : 'Ecris quelque chose sur la qualite' → Precisez le contexte, le format, le public.",
        "Trop long d'un coup : Decoupez les taches complexes en plusieurs prompts successifs.",
        "Pas de format de sortie : Precisez toujours 'en tableau', 'en bullet points', 'en JSON'...",
        "Faire confiance aveuglément : Verifiez toujours les faits critiques, dates, chiffres cites.",
        "Ignorer l'historique : Dans une conversation, chaque message est contexte pour le suivant.",
    ]
    for e in erreurs:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(e)
        r.font.size = Pt(9)

    doc.add_paragraph()
    add_section_title(doc, "Outils recommandes — Formation UIMM-CVDL")
    outils = [
        ("NotebookLM (Google)", "notebooklm.google.com",
         "Interroger ses propres documents PDF/Word. Cree des podcasts audio. Gratuit."),
        ("Antigravity", "Interface locale formation",
         "Agent IA avance. Automatise des taches, ecrit du code, navigue le web."),
        ("Gemini", "gemini.google.com",
         "LLM Google. Integre a Google Workspace (Docs, Sheets, Gmail). Compte UIMM."),
        ("ChatGPT", "chatgpt.com",
         "LLM OpenAI. Interface conviviale. Version gratuite suffisante pour la formation."),
    ]
    tbl4 = doc.add_table(rows=len(outils)+1, cols=3)
    tbl4.style = 'Table Grid'
    for ci, h in enumerate(["Outil", "Acces", "Usage principal"]):
        set_cell_bg(tbl4.cell(0, ci), '1A376C')
        r = tbl4.cell(0, ci).paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.color.rgb = BLANC
        r.font.size = Pt(10)
    for i, (nom, url, usage) in enumerate(outils, 1):
        if i % 2 == 0:
            for ci in range(3):
                set_cell_bg(tbl4.cell(i, ci), 'EEF2F7')
        r1 = tbl4.cell(i, 0).paragraphs[0].add_run(nom)
        r1.font.bold = True
        r1.font.size = Pt(9)
        r1.font.color.rgb = BLEU
        r2 = tbl4.cell(i, 1).paragraphs[0].add_run(url)
        r2.font.size = Pt(8)
        r3 = tbl4.cell(i, 2).paragraphs[0].add_run(usage)
        r3.font.size = Pt(9)

    path = os.path.join(OUT_DIR, 'Fiche_Apprenant_CheatSheet_Prompt.docx')
    doc.save(path)
    print(f"[OK] {path}")


if __name__ == '__main__':
    build_carnet()
    build_cheatsheet()
    print("\nGeneration terminee. Deux fichiers DOCX crees dans 00-Formation/")
