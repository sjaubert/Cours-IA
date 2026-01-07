from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_note():
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('Note de Synthèse : L\'Art du Prompting avec Antigravity', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_paragraph("Cette note a pour objectif d'aider les étudiants en Bachelor à structurer leurs interactions avec l'IA, que ce soit via une interface de chat classique ou directement au sein de l'IDE Antigravity.")

    # Section 1 : Les Fondements d'un Prompt Réussi (Méthode ROLE/CTCF)
    doc.add_heading('1. La Structure ROLE/CTCF', level=1)
    p = doc.add_paragraph()
    p.add_run('Pour obtenir une réponse exploitable en maintenance, structurez votre prompt selon ces 5 piliers :').italic = True
    
    doc.add_paragraph('Rôle : Qui est l\'IA ? (ex: "Tu es un expert en maintenance électromécanique").', style='List Bullet')
    doc.add_paragraph('Contexte : Quel est l\'environnement ? (ex: "Sur une ligne de conditionnement, le moteur M2 s\'arrête de manière aléatoire").', style='List Bullet')
    doc.add_paragraph('Tâche : Que doit faire l\'IA ? (ex: "Analyse ces logs et propose 3 hypothèses de pannes").', style='List Bullet')
    doc.add_paragraph('Contraintes : Quelles sont les limites ? (ex: "Ne propose que des solutions vérifiables sans démontage lourd").', style='List Bullet')
    doc.add_paragraph('Format : Comment doit être la réponse ? (ex: "Réponds sous forme de tableau comparatif").', style='List Bullet')

    # Section 2 : Les Attentes de l'IA (LLM)
    doc.add_heading('2. Ce que l\'IA attend de vous', level=1)
    doc.add_paragraph("L'IA n'est pas un devin technique. Pour éviter les \"hallucinations\" (réponses inventées) :")
    doc.add_paragraph('Précision Technique : Donnez des valeurs exactes (tensions, courants, fréquences).', style='List Number')
    doc.add_paragraph('Itération : Ne vous contentez pas d\'un seul prompt. Affinez en disant "Précise l\'aspect électrique de ta deuxième hypothèse".', style='List Number')
    doc.add_paragraph('Fourniture de données brutes : Copiez-collez des extraits de logs ou de manuels pour "ancrer" l\'IA dans la réalité.', style='List Number')

    # Section 3 : Usage Spécifique en IDE (Antigravity)
    doc.add_heading('3. L\'Assistance dans l\'IDE Antigravity', level=1)
    doc.add_paragraph("L'IDE offre une dimension supplémentaire grâce à la connaissance du code et des données du projet :")
    
    doc.add_paragraph('Utilisation du Contexte : Ouvrez les fichiers pertinents dans vos onglets. Antigravity "voit" votre code et vos fichiers de données (CSV, JSON).', style='List Bullet')
    doc.add_paragraph('Bibliothèques Spécifiques : Utilisez l\'IA pour générer des scripts d\'analyse avec Pandas ou Matplotlib (ex: "Aide-moi à tracer une courbe de courant à partir de @logs_convoyeur.csv").', style='List Bullet')
    doc.add_paragraph('Débogage Guidé : En cas d\'erreur de script de maintenance, demandez "Pourquoi ce calcul de MTBF me donne une erreur de division par zéro ?"', style='List Bullet')

    # Conclusion
    doc.add_paragraph("\nEn résumé : Un bon prompt fait gagner des heures de diagnostic. Un mauvais prompt fait perdre du temps en corrections inutiles.")

    # Sauvegarde
    path = r'c:\Users\s.jaubert\projets\Cours-IA\Formation_IA_Usages_Industrie_UIMM-CVDL\Note_Synthese_Prompting_Antigravity.docx'
    doc.save(path)
    print(f"Fichier créé : {path}")

if __name__ == "__main__":
    create_note()
